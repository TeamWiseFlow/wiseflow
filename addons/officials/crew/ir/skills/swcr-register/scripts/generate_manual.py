#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate software operation manual (DOCX) from README.

Converts a README.md file into a formatted Word document suitable for
software copyright registration "文档鉴别材料" (Document Identification Material).

Features:
- Proper heading hierarchy (H0/H1/H2)
- Code blocks with monospace font
- Bullet lists
- Tables (basic support)
- Header: software name + version
- Chinese font: SimHei; Code font: Courier New

Usage:
    python generate_manual.py \
        --title "智能数据分析平台软件" \
        --version "V1.0" \
        --readme /path/to/README.md \
        --output manual.docx
"""

import argparse
import re
import sys
from typing import List, Tuple

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


def parse_markdown(content: str) -> List[Tuple[str, str]]:
    """Parse markdown into a list of (type, content) tokens.

    Token types:
        h0, h1, h2, h3 - headings
        code_start, code_end - code block delimiters
        code_line - line inside code block
        list_item - bullet list item
        table_row - table row (pipe-separated)
        table_separator - table separator line (---|---)
        paragraph - regular text
        blank - empty line
        hr - horizontal rule
    """
    tokens: List[Tuple[str, str]] = []
    in_code_block = False

    for line in content.split("\n"):
        if line.strip().startswith("```"):
            if in_code_block:
                tokens.append(("code_end", ""))
                in_code_block = False
            else:
                lang = line.strip()[3:].strip()
                tokens.append(("code_start", lang))
                in_code_block = True
            continue

        if in_code_block:
            tokens.append(("code_line", line))
            continue

        stripped = line.strip()

        if not stripped:
            tokens.append(("blank", ""))
        elif stripped == "---" or stripped == "***" or stripped == "___":
            tokens.append(("hr", ""))
        elif line.startswith("# "):
            tokens.append(("h0", stripped[2:]))
        elif line.startswith("## "):
            tokens.append(("h1", stripped[3:]))
        elif line.startswith("### "):
            tokens.append(("h2", stripped[4:]))
        elif line.startswith("#### "):
            tokens.append(("h3", stripped[5:]))
        elif re.match(r"^\|.*\|$", stripped):
            if re.match(r"^\|[\s\-:|]+\|$", stripped):
                tokens.append(("table_separator", stripped))
            else:
                tokens.append(("table_row", stripped))
        elif re.match(r"^[\s]*[-*+]\s", line):
            text = re.sub(r"^[\s]*[-*+]\s", "", line)
            tokens.append(("list_item", text))
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            tokens.append(("list_item", text))
        else:
            tokens.append(("paragraph", stripped))

    return tokens


def add_header(doc: Document, title: str, version: str) -> None:
    """Add document header with software name and version."""
    section = doc.sections[0]
    header = section.header
    header_para = header.paragraphs[0]
    header_para.text = f"{title} {version}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in header_para.runs:
        run.font.size = Pt(9)


def set_chinese_font(doc: Document) -> None:
    """Set default Chinese font to SimHei."""
    style = doc.styles["Normal"]
    style.font.name = "SimHei"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "SimHei")


def add_table_from_rows(doc: Document, rows: List[str]) -> None:
    """Add a simple table from pipe-separated row strings."""
    if not rows:
        return

    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip("|").split("|")]
        parsed.append(cells)

    if not parsed:
        return

    num_cols = max(len(r) for r in parsed)
    table = doc.add_table(rows=len(parsed), cols=num_cols)
    table.style = "Table Grid"

    for i, row_data in enumerate(parsed):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = cell_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)


def build_document(
    tokens: List[Tuple[str, str]],
    title: str,
    version: str,
) -> Document:
    """Build the DOCX document from parsed markdown tokens."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(0.8)

    set_chinese_font(doc)
    add_header(doc, title, version)

    pending_paragraph: List[str] = []
    pending_table_rows: List[str] = []

    def flush_paragraph() -> None:
        if pending_paragraph:
            text = " ".join(pending_paragraph)
            p = doc.add_paragraph(text)
            for run in p.runs:
                run.font.size = Pt(10.5)
            pending_paragraph.clear()

    def flush_table() -> None:
        if pending_table_rows:
            add_table_from_rows(doc, pending_table_rows)
            pending_table_rows.clear()

    for token_type, content in tokens:
        if token_type in ("h0", "h1", "h2", "h3"):
            flush_paragraph()
            flush_table()
            level = int(token_type[1])
            doc.add_heading(content, level=level)

        elif token_type == "code_start":
            flush_paragraph()
            flush_table()

        elif token_type == "code_end":
            flush_paragraph()

        elif token_type == "code_line":
            p = doc.add_paragraph()
            run = p.add_run(content)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0

        elif token_type == "list_item":
            flush_paragraph()
            flush_table()
            p = doc.add_paragraph(content, style="List Bullet")
            for run in p.runs:
                run.font.size = Pt(10.5)

        elif token_type == "table_row":
            flush_paragraph()
            pending_table_rows.append(content)

        elif token_type == "table_separator":
            pass  # skip separator lines

        elif token_type == "hr":
            flush_paragraph()
            flush_table()
            p = doc.add_paragraph()
            p.add_run("—" * 40)

        elif token_type == "blank":
            flush_paragraph()
            flush_table()

        elif token_type == "paragraph":
            flush_table()
            pending_paragraph.append(content)

    flush_paragraph()
    flush_table()

    return doc


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate software operation manual (DOCX) from README."
    )
    parser.add_argument("--title", required=True, help="Software full name")
    parser.add_argument("--version", default="V1.0", help="Software version")
    parser.add_argument("--readme", required=True, help="Path to README.md")
    parser.add_argument("--output", required=True, help="Output DOCX file path")

    args = parser.parse_args()

    if not DOCX_AVAILABLE:
        print("Error: python-docx is required. Install with: pip install python-docx")
        return 1

    try:
        with open(args.readme, "r", encoding="utf-8") as f:
            content = f.read()
    except FileNotFoundError:
        print(f"Error: README file not found: {args.readme}")
        return 1
    except Exception as exc:
        print(f"Error reading README: {exc}")
        return 1

    tokens = parse_markdown(content)
    doc = build_document(tokens, args.title, args.version)
    doc.save(args.output)

    print(f"Operation manual created: {args.output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
