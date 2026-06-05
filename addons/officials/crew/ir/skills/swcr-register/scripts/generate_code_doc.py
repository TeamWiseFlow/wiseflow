#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate software copyright registration source code document (DOCX).

Produces a formatted Word document containing source code that meets
Chinese software copyright authority requirements:
- Each page contains 50 effective lines of code (non-blank, non-comment)
- Front 30 pages + back 30 pages with an ellipsis page in between
- Header: software name + version (left) + page number (right)
- Code font: Courier New 8pt; Chinese auxiliary font: SimSun

Usage:
    python generate_code_doc.py \
        --title "智能数据分析平台软件" \
        --version "V1.0" \
        --source-dir /path/to/code \
        --output output.docx
"""

import argparse
import codecs
import logging
import sys
from os import scandir
from os.path import abspath, relpath
from typing import List

try:
    import chardet
    CHARDET_AVAILABLE = True
except ImportError:
    CHARDET_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

logger = logging.getLogger(__name__)

DEFAULT_EXTS = [
    "c", "h", "py", "js", "ts", "java", "cpp", "hpp",
    "go", "rs", "rb", "php", "cs", "swift", "kt", "scala",
    "jsx", "tsx", "vue", "svelte",
]
DEFAULT_COMMENT_CHARS = ["/*", "* ", "*/", "//", "#"]
DEFAULT_EXCLUDES = [
    "node_modules", ".git", "__pycache__", "venv", ".venv",
    "dist", "build", ".next", ".nuxt", "target", "bin",
    "obj", ".idea", ".vscode", "coverage", ".cache",
]


def detect_encoding(file_path: str) -> str:
    """Detect file encoding using chardet with fallback."""
    if not CHARDET_AVAILABLE:
        return "utf-8"

    with open(file_path, "rb") as fd:
        raw_data = fd.read(32768)  # 32KB sample is sufficient for chardet

    result = chardet.detect(raw_data)
    encoding = result.get("encoding", "utf-8")
    confidence = result.get("confidence", 0)

    if confidence < 0.7:
        for enc in ["utf-8", "gbk", "gb2312", "big5", "latin-1"]:
            try:
                raw_data.decode(enc)
                encoding = enc
                break
            except (UnicodeDecodeError, LookupError):
                continue

    return encoding


def find_code_files(
    source_dir: str,
    exts: List[str],
    excludes: List[str],
) -> List[str]:
    """Recursively find source code files matching extensions, excluding dirs."""
    files: List[str] = []
    abs_excludes = [abspath(e) for e in excludes]

    def _should_exclude(path: str) -> bool:
        abs_path = abspath(path)
        return any(abs_path.startswith(ex) for ex in abs_excludes)

    def _scan(directory: str) -> None:
        try:
            for entry in scandir(directory):
                if entry.name.startswith("."):
                    continue
                if _should_exclude(entry.path):
                    continue
                if entry.is_file():
                    if any(entry.name.endswith(f".{ext}") for ext in exts):
                        files.append(abspath(entry.path))
                elif entry.is_dir():
                    _scan(entry.path)
        except PermissionError:
            logger.warning("Permission denied: %s", directory)

    _scan(source_dir)
    return files


def is_blank_line(line: str) -> bool:
    return not bool(line.strip())


def is_comment_line(line: str, comment_chars: List[str]) -> bool:
    stripped = line.lstrip()
    return any(stripped.startswith(cc) for cc in comment_chars)


def wrap_long_line(line: str, max_chars: int = 90) -> List[str]:
    if len(line) <= max_chars:
        return [line]
    wrapped = []
    while len(line) > max_chars:
        wrapped.append(line[:max_chars])
        line = line[max_chars:]
    if line:
        wrapped.append(line)
    return wrapped


def collect_code_lines(
    files: List[str],
    comment_chars: List[str],
    base_dir: str,
) -> List[str]:
    """Read all source files and collect code lines."""
    all_lines: List[str] = []

    for filepath in files:
        encoding = detect_encoding(filepath)
        logger.info("Processing: %s (encoding: %s)", filepath, encoding)

        try:
            relative = relpath(filepath, base_dir)
        except ValueError:
            relative = filepath

        all_lines.append(f"# File: {relative}")

        try:
            with codecs.open(filepath, "r", encoding, errors="replace") as fp:
                for line in fp:
                    line = line.rstrip()
                    all_lines.extend(wrap_long_line(line, max_chars=90))
        except Exception as exc:
            logger.error("Error reading %s: %s", filepath, exc)
            all_lines.append(f"# Error reading file: {exc}")

    return all_lines


def count_effective_lines(lines: List[str], comment_chars: List[str]) -> int:
    """Count non-blank, non-comment lines."""
    return sum(
        1 for line in lines
        if not is_blank_line(line) and not is_comment_line(line, comment_chars)
    )


def _is_effective(line: str, comment_chars: List[str]) -> bool:
    """Check if a line is effective (non-blank and non-comment)."""
    return not is_blank_line(line) and not is_comment_line(line, comment_chars)


def split_into_pages(
    all_lines: List[str],
    comment_chars: List[str],
    lines_per_page: int = 50,
    max_front_pages: int = 30,
    max_back_pages: int = 30,
) -> tuple:
    """Split code lines into front pages and back pages.

    Each page contains lines_per_page effective lines (non-blank, non-comment).
    """
    if not all_lines:
        return [], []

    front_pages: List[List[str]] = []
    back_pages: List[List[str]] = []

    current_page: List[str] = []
    effective_count = 0
    page_count = 0

    i = 0
    while i < len(all_lines) and page_count < max_front_pages:
        line = all_lines[i]
        current_page.append(line)
        if _is_effective(line, comment_chars):
            effective_count += 1
        if effective_count >= lines_per_page or i == len(all_lines) - 1:
            front_pages.append(current_page.copy())
            logger.info(
                "Front page %d: %d lines, %d effective",
                page_count + 1, len(current_page), effective_count,
            )
            current_page = []
            effective_count = 0
            page_count += 1
        i += 1

    if i < len(all_lines):
        remaining = all_lines[i:]
        remaining_effective = count_effective_lines(remaining, comment_chars)

        if remaining_effective > max_back_pages * lines_per_page:
            target = max_back_pages * lines_per_page
            start_pos = len(remaining) - 1
            eff = 0
            for j in range(len(remaining) - 1, -1, -1):
                if _is_effective(remaining[j], comment_chars):
                    eff += 1
                if eff >= target:
                    start_pos = j
                    break
            back_source = remaining[start_pos:]
        else:
            back_source = remaining

        current_page = []
        effective_count = 0
        for line in back_source:
            current_page.append(line)
            if _is_effective(line, comment_chars):
                effective_count += 1
            if effective_count >= lines_per_page:
                back_pages.append(current_page.copy())
                current_page = []
                effective_count = 0
        if current_page:
            back_pages.append(current_page)

    return front_pages, back_pages


def create_docx(
    filename: str,
    title: str,
    version: str,
    front_pages: List[List[str]],
    back_pages: List[List[str]],
) -> None:
    """Create the source code DOCX document."""
    if not DOCX_AVAILABLE:
        print("Error: python-docx is required. Install with: pip install python-docx")
        sys.exit(1)

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.5)

    total_front = len(front_pages)

    for page_num, page_lines in enumerate(front_pages, 1):
        _add_code_page(doc, page_lines, page_num, title, version)
        if page_num < total_front or back_pages:
            doc.add_page_break()

    if back_pages:
        _add_ellipsis_page(doc, total_front + 1, title, version)
        for page_num, page_lines in enumerate(back_pages, total_front + 2):
            doc.add_page_break()
            _add_code_page(doc, page_lines, page_num, title, version)

    doc.save(filename)


def _add_code_page(
    doc: Document,
    lines: List[str],
    page_num: int,
    title: str,
    version: str,
) -> None:
    """Add one page of code to the document."""
    header_text = f"{title} {version}"
    p = doc.add_paragraph()
    run = p.add_run(header_text)
    run.font.size = Pt(9)
    run = p.add_run(f"\t\t\t\t\t\t\t\t\t\t{page_num}")
    run.font.size = Pt(9)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.0

    p = doc.add_paragraph()
    run = p.add_run("_" * 95)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0

    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line if line.strip() else " ")
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = run._element
        r.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")


def _add_ellipsis_page(
    doc: Document,
    page_num: int,
    title: str,
    version: str,
) -> None:
    """Add an ellipsis page to the document."""
    header = f"{title} {version}"
    p = doc.add_paragraph()
    run = p.add_run(header)
    run.font.size = Pt(10)
    run = p.add_run(f"\t\t\t\t\t\t\t\t\t\t{page_num}")
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.add_run("_" * 100)

    for _ in range(20):
        doc.add_paragraph()

    p = doc.add_paragraph()
    run = p.add_run("......")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run.font.size = Pt(24)


def main() -> int:
    parser = argparse.ArgumentParser(
        description="Generate source code document for software copyright registration."
    )
    parser.add_argument("--title", required=True, help="Software full name")
    parser.add_argument("--version", default="V1.0", help="Software version")
    parser.add_argument("--source-dir", required=True, help="Source code directory")
    parser.add_argument("--output", required=True, help="Output DOCX file path")
    parser.add_argument(
        "--exts", nargs="+", default=None,
        help="Source code file extensions (auto-detect if omitted)",
    )
    parser.add_argument(
        "--comment-chars", nargs="+", default=None,
        help="Comment characters (auto-detect if omitted)",
    )
    parser.add_argument(
        "--excludes", nargs="+", default=DEFAULT_EXCLUDES,
        help="Directories to exclude",
    )
    parser.add_argument(
        "--max-front-pages", type=int, default=30,
        help="Max front pages (default: 30)",
    )
    parser.add_argument(
        "--max-back-pages", type=int, default=30,
        help="Max back pages (default: 30)",
    )
    parser.add_argument(
        "--verbose", action="store_true",
        help="Enable verbose logging",
    )

    args = parser.parse_args()

    if args.verbose:
        logging.basicConfig(level=logging.DEBUG)
    else:
        logging.basicConfig(level=logging.INFO)

    if not DOCX_AVAILABLE:
        print("Error: python-docx is required. Install with: pip install python-docx")
        return 1

    exts = args.exts if args.exts else DEFAULT_EXTS
    comment_chars = args.comment_chars if args.comment_chars else DEFAULT_COMMENT_CHARS

    source_dir = abspath(args.source_dir)

    print(f"Scanning source code in: {source_dir}")
    print(f"Extensions: {exts}")
    print(f"Comment chars: {comment_chars}")
    print(f"Excludes: {args.excludes}")

    files = find_code_files(source_dir, exts, args.excludes)
    print(f"Found {len(files)} source code files")

    if not files:
        print("Error: No source code files found. Check --source-dir and --exts.")
        return 1

    all_lines = collect_code_lines(files, comment_chars, source_dir)
    print(f"Total lines collected: {len(all_lines)}")

    effective = count_effective_lines(all_lines, comment_chars)
    print(f"Effective lines (non-blank, non-comment): {effective}")

    front_pages, back_pages = split_into_pages(
        all_lines, comment_chars,
        lines_per_page=50,
        max_front_pages=args.max_front_pages,
        max_back_pages=args.max_back_pages,
    )
    print(f"Front pages: {len(front_pages)}")
    print(f"Back pages: {len(back_pages)}")
    total_pages = len(front_pages) + (1 if back_pages else 0) + len(back_pages)
    print(f"Total pages (including ellipsis): {total_pages}")

    create_docx(args.output, args.title, args.version, front_pages, back_pages)
    print(f"Source code document created: {args.output}")

    return 0


if __name__ == "__main__":
    sys.exit(main())
