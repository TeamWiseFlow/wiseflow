#!/usr/bin/env python3
"""
md_to_docx.py — Convert Markdown to DOCX with embedded local images.

Usage:
    python md_to_docx.py -f article.md -o /tmp/article.docx

Rules:
  - Local images are embedded into the Word document.
  - If the resulting docx exceeds 15 MB, images are stripped and the file is saved again.
  - Remote images (http/https) are skipped with a placeholder.
"""

import argparse
import os
import re
import shutil
import sys
import tempfile
import zipfile
from pathlib import Path


def parse_frontmatter(text: str) -> tuple[dict, str]:
    """Return (meta dict, body text without frontmatter)."""
    meta: dict = {}
    fm_match = re.match(r"^---\n(.*?)\n---\n", text, re.DOTALL)
    if not fm_match:
        return meta, text
    for line in fm_match.group(1).splitlines():
        kv = re.match(r"^(\w+):\s*(.+)", line)
        if kv:
            meta[kv.group(1)] = kv.group(2).strip().strip("\"'")
    return meta, text[fm_match.end():]


def add_inline_runs(paragraph, text: str, base_dir: Path) -> None:
    """Add text runs with bold/italic/inline-code to a paragraph.
    Inline images inside a paragraph are appended as separate runs."""
    from docx.shared import Pt

    # Strip inline images (can't embed mid-paragraph)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"[\1]", text)
    # Hyperlinks: preserve URL as "label (url)"
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)

    # Split on bold/italic markers (greedy-safe patterns)
    token_re = re.compile(
        r"(\*\*\*[^*]+?\*\*\*"
        r"|\*\*[^*]+?\*\*"
        r"|__[^_]+?__"
        r"|\*[^*]+?\*"
        r"|_[^_]+?_"
        r"|`[^`]+?`)"
    )
    pos = 0
    for m in token_re.finditer(text):
        # plain text before match
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("***") or token.startswith("___"):
            run = paragraph.add_run(token[3:-3])
            run.bold = True
            run.italic = True
        elif token.startswith("**") or token.startswith("__"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("*") or token.startswith("_"):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        pos = m.end()
    # remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def try_add_image(doc, img_path: Path, width_inches: float = 5.5) -> bool:
    """Add image paragraph to doc. Returns True on success."""
    from docx.shared import Inches

    if not img_path.exists():
        doc.add_paragraph(f"[图片未找到: {img_path.name}]")
        return False
    try:
        doc.add_picture(str(img_path), width=Inches(width_inches))
        return True
    except Exception:
        # Fallback: convert via Pillow (handles progressive JPEG, RGBA, etc.)
        try:
            import os
            import tempfile
            from PIL import Image as PILImage

            with PILImage.open(img_path) as im:
                rgb = im.convert("RGB")
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as fh:
                    tmp = fh.name
                rgb.save(tmp, "JPEG", quality=90)
            try:
                doc.add_picture(tmp, width=Inches(width_inches))
                return True
            finally:
                os.unlink(tmp)
        except Exception as exc2:
            doc.add_paragraph(f"[图片嵌入失败: {img_path.name} — {exc2}]")
            return False


def convert(md_path: Path, out_path: Path) -> None:
    from docx import Document
    from docx.shared import Pt

    text = md_path.read_text(encoding="utf-8")
    base_dir = md_path.parent
    meta, body = parse_frontmatter(text)

    doc = Document()

    # Title from frontmatter
    if meta.get("title"):
        doc.add_heading(meta["title"], level=0)

    lines = body.splitlines()
    i = 0
    in_code_block = False
    code_lines: list[str] = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        # ── Code block ────────────────────────────────────────────────────────
        if line.startswith("```"):
            if not in_code_block:
                in_code_block = True
                code_lang = line[3:].strip()
                code_lines = []
            else:
                in_code_block = False
                p = doc.add_paragraph("\n".join(code_lines), style="No Spacing")
                if p.runs:
                    p.runs[0].font.name = "Courier New"
                    p.runs[0].font.size = Pt(10)
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # ── Heading ───────────────────────────────────────────────────────────
        h_match = re.match(r"^(#{1,6})\s+(.+)", line)
        if h_match:
            doc.add_heading(h_match.group(2).strip(), level=len(h_match.group(1)))
            i += 1
            continue

        # ── Standalone image (whole line) ─────────────────────────────────────
        img_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$", line.strip())
        if img_match:
            src = img_match.group(2)
            if not src.startswith("http"):
                img_path = base_dir / src
                if not img_path.exists():
                    img_path = Path.cwd() / src  # fallback: workspace root
                try_add_image(doc, img_path)
            else:
                doc.add_paragraph(f"[远程图片: {src}]")
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────────
        if re.match(r"^[-*_]{3,}\s*$", line):
            doc.add_paragraph("─" * 40)
            i += 1
            continue

        # ── Unordered list ────────────────────────────────────────────────────
        ul_match = re.match(r"^[\-\*\+]\s+(.+)", line)
        if ul_match:
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, ul_match.group(1), base_dir)
            i += 1
            continue

        # ── Ordered list ──────────────────────────────────────────────────────
        ol_match = re.match(r"^\d+\.\s+(.+)", line)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, ol_match.group(1), base_dir)
            i += 1
            continue

        # ── Blockquote ────────────────────────────────────────────────────────
        bq_match = re.match(r"^>\s+(.*)", line)
        if bq_match:
            p = doc.add_paragraph(style="Quote")
            add_inline_runs(p, bq_match.group(1), base_dir)
            i += 1
            continue

        # ── Table ─────────────────────────────────────────────────────────────
        if "|" in line and re.match(r"^\s*\|", line):
            # Peek ahead to confirm next non-empty line is a separator row
            j = i + 1
            while j < len(lines) and not lines[j].strip():
                j += 1
            if j < len(lines) and re.match(r"^\s*\|[\s|:\-]+\|?\s*$", lines[j]):
                rows_data: list[list[str]] = []
                while i < len(lines):
                    row_line = lines[i]
                    if not row_line.strip() or not re.match(r"^\s*\|", row_line):
                        break
                    if re.match(r"^\s*\|[\s|:\-]+\|?\s*$", row_line):
                        i += 1
                        continue  # skip separator row
                    cells = [c.strip() for c in row_line.strip().strip("|").split("|")]
                    rows_data.append(cells)
                    i += 1
                if rows_data:
                    num_cols = max(len(r) for r in rows_data)
                    tbl = doc.add_table(rows=len(rows_data), cols=num_cols)
                    tbl.style = "Table Grid"
                    for r_idx, row_cells in enumerate(rows_data):
                        for c_idx in range(num_cols):
                            cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ""
                            cell = tbl.cell(r_idx, c_idx)
                            cell.text = ""
                            para = cell.paragraphs[0]
                            add_inline_runs(para, cell_text, base_dir)
                            if r_idx == 0:
                                for run in para.runs:
                                    run.bold = True
                continue
            # not a table — fall through to normal paragraph

        # ── Empty line ────────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Normal paragraph ──────────────────────────────────────────────────
        p = doc.add_paragraph()
        add_inline_runs(p, line, base_dir)
        i += 1

    doc.save(str(out_path))


def strip_images_from_docx(docx_path: Path) -> None:
    """Remove all embedded images from a docx to reduce file size."""
    tmp = docx_path.with_suffix(".tmp.docx")
    shutil.copy2(docx_path, tmp)

    # Step 1: rebuild zip without word/media/* files
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as fh:
        stage1 = Path(fh.name)

    with zipfile.ZipFile(tmp, "r") as zin, zipfile.ZipFile(
        stage1, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            if item.filename.startswith("word/media/"):
                continue
            zout.writestr(item, zin.read(item.filename))

    # Step 2: strip <w:drawing>…</w:drawing> blocks from document.xml
    with zipfile.ZipFile(stage1, "r") as z:
        doc_xml = z.read("word/document.xml").decode("utf-8")
    doc_xml = re.sub(r"<w:drawing>.*?</w:drawing>", "", doc_xml, flags=re.DOTALL)

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as fh:
        stage2 = Path(fh.name)

    with zipfile.ZipFile(stage1, "r") as zin, zipfile.ZipFile(
        stage2, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            if item.filename == "word/document.xml":
                zout.writestr(item, doc_xml.encode("utf-8"))
            else:
                zout.writestr(item, zin.read(item.filename))

    shutil.move(str(stage2), str(docx_path))
    tmp.unlink(missing_ok=True)
    stage1.unlink(missing_ok=True)


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert Markdown to DOCX")
    parser.add_argument("-f", "--file", required=True, help="Input Markdown file")
    parser.add_argument("-o", "--output", help="Output DOCX path (default: same dir as input)")
    args = parser.parse_args()

    md_path = Path(args.file).resolve()
    out_path = Path(args.output).resolve() if args.output else md_path.with_suffix(".docx")

    if not md_path.exists():
        print(f"ERROR: 文件不存在: {md_path}", file=sys.stderr)
        return 1

    try:
        from docx import Document  # noqa: F401 — early import check
    except ImportError:
        print(
            "ERROR: 缺少依赖 python-docx。请运行：pip install python-docx",
            file=sys.stderr,
        )
        return 1

    print(f">>> 转换: {md_path.name} → {out_path.name}")
    convert(md_path, out_path)

    size_mb = out_path.stat().st_size / (1024 * 1024)
    print(f">>> 文件大小: {size_mb:.1f} MB")

    if size_mb > 15:
        print(">>> 超过 15 MB，移除图片后重新保存...")
        strip_images_from_docx(out_path)
        size_mb = out_path.stat().st_size / (1024 * 1024)
        print(f">>> 移除图片后大小: {size_mb:.1f} MB（图片已删除，请在头条编辑器中手动补图）")

    print(f">>> 完成: {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
