import random
import re
from llms.dashscope_wrapper import dashscope_llm
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from general_utils import isChinesePunctuation
import configparser

# qwen-72b-chat支持最大30k输入，考虑prompt其他部分，content不应超过30000字符长度
# 如果换qwen-max（最大输入6k),这里就要换成6000,但这样很多文章不能分析了
# 本地部署模型（qwen-14b这里可能仅支持4k输入，可能根本这套模式就行不通）
max_input_tokens = 30000
config = configparser.ConfigParser()
config.read('../config.ini')


def get_report(insigt: str, articles: list[dict], memory: str, topics: list[str], comment: str, docx_file: str, logger=None) -> (bool, str):
    zh_index = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十', '十一', '十二']

    if isChinesePunctuation(insigt[-1]):
        insigt = insigt[:-1]

    # 分离段落和标题
    if len(topics) == 0:
        title = ''
    elif len(topics) == 1:
        title = topics[0]
        topics = []
    else:
        title = topics[0]
        topics = [s.strip() for s in topics[1:] if s.strip()]

    schema = f'【标题】{title}\n\n【综述】\n\n'
    if topics:
        for i in range(len(topics)):
            schema += f'【{zh_index[i]}、{topics[i]}】\n\n'

    # 先判断是否是修改要求（有原文和评论，且原文的段落要求与给到的topics一致）
    system_prompt, user_prompt = '', ''
    if memory and comment:
        paragraphs = re.findall("、(.*?)】", memory)
        if set(topics) <= set(paragraphs):
            logger.debug("no change in Topics, need modified the report")
            system_prompt = f'''你是一名{config['prompts']['character']}，你近日向上级提交了一份{config['prompts']['report_type']}报告，如下是报告原文。接下来你将收到来自上级部门的修改意见，请据此修改你的报告：
报告原文： 
"""{memory}"""
'''
            user_prompt = f'上级部门修改意见："""{comment}"""'

    if not system_prompt or not user_prompt:
        logger.debug("need generate the report")
        texts = ''
        for article in articles:
            if article['content']:
                texts += f"<article>{article['content']}</article>\n"
            else:
                if article['abstract']:
                    texts += f"<article>{article['abstract']}</article>\n"
                else:
                    texts += f"<article>{article['title']}</article>\n"

            if len(texts) > max_input_tokens:
                break

        logger.debug(f"articles context length: {len(texts)}")
        system_prompt = f'''你是一名{config['prompts']['character']}，在近期的工作中我们从所关注的网站中发现了一条重要的{config['prompts']['report_type']}线索，线索和相关文章（用XML标签分隔）如下：
情报线索： """{insigt} """
相关文章：
{texts}
现在请基于这些信息按要求输出专业的书面报告。'''

        if comment:
            user_prompt = (f'1、不管原始资料是什么语言，你必须使用简体中文输出报告，除非是人名、组织和机构的名称、缩写；'
                           f'2、对事实的陈述务必基于所提供的相关文章，绝对不可以臆想；3、{comment}。\n')
        else:
            user_prompt = ('1、不管原始资料是什么语言，你必须使用简体中文输出报告，除非是人名、组织和机构的名称、缩写；'
                           '2、对事实的陈述务必基于所提供的相关文章，绝对不可以臆想。')

    user_prompt += f'\n请按如下格式输出你的报告：\n{schema}'

    # 生成阶段
    check_flag = False
    check_list = schema.split('\n\n')
    check_list = [_[1:] for _ in check_list if _.startswith('【')]
    result = ''
    for i in range(2):
        result = dashscope_llm([{'role': 'system', 'content': system_prompt}, {'role': 'user', 'content': user_prompt}],
                               'qwen1.5-72b-chat', seed=random.randint(1, 10000), logger=logger)
        logger.debug(f"raw result:\n{result}")
        if len(result) > 50:
            check_flag = True
            for check_item in check_list[2:]:
                if check_item not in result:
                    check_flag = False
                    break
        if check_flag:
            break

        logger.debug("result not good, re-generating...")

    if not check_flag:
        # 这里其实存在两种情况，一个是llm失效，一个是多次尝试后生成结果还是不行
        if not result:
            logger.warning('report-process-error: LLM out of work!')
            return False, ''
        else:
            logger.warning('report-process-error: cannot generate, change topics and insight, then re-try')
            return False, ''

    # parse process
    contents = result.split("【")
    bodies = {}
    for text in contents:
        for item in check_list:
            if text.startswith(item):
                check_list.remove(item)
                key, value = text.split("】")
                value = value.strip()
                if isChinesePunctuation(value[0]):
                    value = value[1:]
                bodies[key] = value.strip()
                break

    if not bodies:
        logger.warning('report-process-error: cannot generate, change topics and insight, then re-try')
        return False, ''

    if '标题' not in bodies:
        if "】" in contents[0]:
            _title = contents[0].split("】")[0]
            bodies['标题'] = _title.strip()
        else:
            if len(contents) > 1 and "】" in contents[1]:
                _title = contents[0].split("】")[0]
                bodies['标题'] = _title.strip()
            else:
                bodies['标题'] = ""

    doc = Document()
    doc.styles['Normal'].font.name = u'宋体'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    doc.styles['Normal'].font.size = Pt(12)
    doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)

    # 先写好标题和摘要
    if not title:
        title = bodies['标题']

    Head = doc.add_heading(level=1)
    Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = Head.add_run(title)
    run.font.name = u'Cambria'
    run.font.color.rgb = RGBColor(0, 0, 0)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Cambria')

    doc.add_paragraph(
        f"\n生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    del bodies['标题']
    if '综述' in bodies:
        doc.add_paragraph(f"\t{bodies['综述']}\n")
        del bodies['综述']

    # 逐段添加章节
    for key, value in bodies.items():
        Head = doc.add_heading(level=2)
        run = Head.add_run(key)
        run.font.name = u'Cambria'
        run.font.color.rgb = RGBColor(0, 0, 0)
        doc.add_paragraph(f"{value}\n")

    # 添加附件引用信息源
    Head = doc.add_heading(level=2)
    run = Head.add_run("附：原始信息网页")
    run.font.name = u'Cambria'
    run.font.color.rgb = RGBColor(0, 0, 0)

    contents = []
    for i, article in enumerate(articles):
        date_text = str(article['publish_time'])
        if len(date_text) == 8:
            date_text = f"{date_text[:4]}-{date_text[4:6]}-{date_text[6:]}"

        contents.append(f"{i+1}、{article['title']}|{date_text}\n{article['url']} ")

    doc.add_paragraph("\n\n".join(contents))

    doc.save(docx_file)

    return True, result[result.find("【"):]
